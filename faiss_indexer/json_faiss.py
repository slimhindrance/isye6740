import os
import json
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import faiss
from dotenv import load_dotenv
from tqdm import tqdm
from docx import Document as DocxDocument  # For DOCX processing
import fitz  # PyMuPDF for PDF processing

load_dotenv()

# Define paths
faiss_index_path = os.getenv("FAISS_INDEX_PATH")#, "/app/faiss")

json_folder = os.getenv("JSON_THREADS_DIR")#, "/app/json_threads")
docx_folder = os.getenv("DOCX_DIR")#, "/app/json_threads")
pdf_folder = os.getenv("PDF_DIR")#, "/app/json_threads")
txt_folder = os.getenv("TXT_DIR")


# Load the embedding model
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")

# Always create a new FAISS index
print("🆕 Creating a new FAISS index...")
os.makedirs(faiss_index_path, exist_ok=True)  # Ensure save directory exists

# Initialize FAISS index
index = faiss.IndexFlatL2(len(embeddings.embed_query("sample")))  # ✅ Use len() instead of .shape[0]

vector_store = FAISS(
    embedding_function=embeddings,
    index=index,
    docstore=InMemoryDocstore(),
    index_to_docstore_id={},
)

# Initialize the text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,  # Adjust the size depending on your needs
    chunk_overlap=200,  # Overlap between chunks for better context continuity
)

# Process JSON, DOCX, and PDF files
json_documents = []
for filename in tqdm(os.listdir(json_folder), desc="Processing JSON Files", unit="file"):

    # Processing JSON Threads
    if filename.endswith(".json"):
        with open(os.path.join(json_folder, filename), 'r', encoding='utf-8') as file:
            thread_data = json.load(file)
            
            # Extract thread content
            thread_id = thread_data.get('id')
            thread_content = f"Thread ID: {thread_id}\nUser {thread_data['user_id']} (Parent Post) says:\n{thread_data.get('document', '')}\n"

            # Add comments and answers
            for comment in thread_data.get('comments', []):
                thread_content += f"Comment from User {comment['user_id']}:\n{comment.get('document', '')}\n"
                for sub_comment in comment.get('comments', []):
                    thread_content += f"  Reply from User {sub_comment['user_id']}:\n  {sub_comment.get('document', '')}\n"

            for answer in thread_data.get('answers', []):
                thread_content += f"Answer from User {answer['user_id']}:\n{answer.get('document', '')}\n"
                for sub_comment in answer.get('comments', []):
                    thread_content += f"  Comment from User {sub_comment['user_id']}:\n  {sub_comment.get('document', '')}\n"

            # Add JSON threads with medium authority
            json_documents.append(Document(page_content=thread_content, metadata={"source": filename, "authority_level": 1}))

docx_documents = []
for filename in tqdm(os.listdir(docx_folder), desc="Processing DOCX Files", unit="file"):
    # Processing DOCX Files
    if filename.endswith(".docx"):
        docx_path = os.path.join(docx_folder, filename)
        docx = DocxDocument(docx_path)

        # Extract text from all paragraphs
        docx_text = "\n".join([para.text for para in docx.paragraphs if para.text.strip() != ""])

        # Split DOCX text into manageable chunks
        chunks = text_splitter.split_text(docx_text)
        for chunk in chunks:
            docx_documents.append(Document(page_content=chunk, metadata={"source": filename, "authority_level": 3}))

pdf_documents = []
for filename in tqdm(os.listdir(pdf_folder), desc="Processing PDF Files", unit="file"):
    # Processing PDF Files
    if filename.endswith(".pdf"):
        pdf_path = os.path.join(pdf_folder, filename)
        pdf_doc = fitz.open(pdf_path)

        pdf_text = ""
        for page in pdf_doc:
            pdf_text += page.get_text()

        pdf_doc.close()

        # Split PDF text into chunks
        chunks = text_splitter.split_text(pdf_text.strip())
        for chunk in chunks:
            pdf_documents.append(Document(page_content=chunk, metadata={"source": filename, "authority_level": 3}))

txt_documents = []
for filename in tqdm(os.listdir(txt_folder), desc="Processing TXT Files", unit="file"):
    if filename.endswith(".txt"):
        txt_path = os.path.join(txt_folder, filename)

        # Open and read the TXT file
        with open(txt_path, 'r', encoding='utf-8') as file:
            txt_text = file.read()

        # Split TXT text into manageable chunks
        chunks = text_splitter.split_text(txt_text)

        # Append each chunk to the document list with metadata
        for chunk in chunks:
            txt_documents.append(Document(page_content=chunk, metadata={"source": filename, "authority_level": 3}))

# Add documents to FAISS and save index
if json_documents:
    vector_store.add_documents(json_documents)
    vector_store.save_local(faiss_index_path)
    print(f"✅ FAISS index updated with {len(json_documents)} JSON documents!")

# Add documents to FAISS and save index
if docx_documents:
    vector_store.add_documents(docx_documents)
    vector_store.save_local(faiss_index_path)
    print(f"✅ FAISS index updated with {len(docx_documents)} DOCX documents!")

# Add documents to FAISS and save index
if pdf_documents:
    vector_store.add_documents(pdf_documents)
    vector_store.save_local(faiss_index_path)
    print(f"✅ FAISS index updated with {len(pdf_documents)} PDF documents!")

# Add documents to FAISS and save index
if txt_documents:
    vector_store.add_documents(txt_documents)
    vector_store.save_local(faiss_index_path)
    print(f"✅ FAISS index updated with {len(txt_documents)} TXT documents!")

else:
    print("❌ No documents found to add.")